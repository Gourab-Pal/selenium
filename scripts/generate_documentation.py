"""Generate exhaustive DOCX documentation for the Selenium E2E framework."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def add_numbered(doc, text):
    doc.add_paragraph(text, style="List Number")


def add_code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "D9E2F3")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = value
    doc.add_paragraph()
    return table


def build_document():
    doc = Document()

    # Title page
    add_title(doc, "Selenium 4 E2E Automation Framework")
    add_para(doc, "Complete Repository Documentation", bold=True).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    add_para(doc, "Repository: https://github.com/Gourab-Pal/selenium", bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Author: Gourab Pal", bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Generated: June 2026", bold=False).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of Contents placeholder
    add_heading(doc, "Table of Contents", 1)
    toc_items = [
        "1. Project Overview",
        "2. Technology Stack",
        "3. Repository Structure",
        "4. Architecture & Component Connections",
        "5. Configuration Reference",
        "6. Running Tests (Headed & Headless)",
        "7. Allure Reporting",
        "8. CI/CD Pipeline (GitHub Actions)",
        "9. How to Trigger CI/CD",
        "10. Test Suite Details",
        "11. Page Object Model",
        "12. Adding New Tests",
        "13. Troubleshooting",
        "14. Quick Command Reference",
    ]
    for item in toc_items:
        add_bullet(doc, item)
    doc.add_page_break()

    # 1. Project Overview
    add_heading(doc, "1. Project Overview", 1)
    add_para(doc, (
        "This repository is a production-style end-to-end (E2E) web automation framework built with "
        "Selenium 4, TestNG, Gradle, and Allure reporting. It follows the Page Object Model (POM) "
        "design pattern and includes GitHub Actions CI/CD integration."
    ))
    add_para(doc, "Key capabilities:", bold=True)
    add_bullet(doc, "Automated browser testing against public demo sites")
    add_bullet(doc, "Configurable browser (Chrome, Firefox, Edge) and headless/headed execution")
    add_bullet(doc, "Parallel test execution via TestNG")
    add_bullet(doc, "Rich Allure HTML reports with steps, screenshots on failure, and metadata")
    add_bullet(doc, "Automated CI pipeline on every push and pull request")
    add_bullet(doc, "Selenium Manager for automatic browser driver resolution (no manual driver setup)")

    # 2. Technology Stack
    add_heading(doc, "2. Technology Stack", 1)
    add_table(doc, ["Component", "Version", "Purpose"], [
        ["Java", "17 (LTS)", "Programming language"],
        ["Gradle", "9.3.0", "Build and dependency management"],
        ["Selenium", "4.44.0", "Browser automation (latest stable Selenium 4)"],
        ["TestNG", "7.10.2", "Test framework and suite orchestration"],
        ["Allure Adapter", "2.35.2", "TestNG integration for reporting"],
        ["Allure CLI", "2.42.0", "HTML report generation"],
        ["Allure Gradle Plugin", "4.0.2", "Gradle integration (Gradle 9 compatible)"],
        ["AssertJ", "3.27.3", "Fluent test assertions"],
        ["SLF4J Simple", "2.0.17", "Logging for Selenium"],
        ["GitHub Actions", "N/A", "CI/CD pipeline"],
    ])

    # 3. Repository Structure
    add_heading(doc, "3. Repository Structure", 1)
    add_code_block(doc, """selenium/
├── .github/
│   └── workflows/
│       └── e2e-tests.yml          # GitHub Actions CI/CD pipeline
├── gradle/
│   └── wrapper/                   # Gradle wrapper (ensures consistent Gradle version)
├── scripts/
│   └── generate_documentation.py  # This documentation generator
├── src/
│   └── test/
│       ├── java/org/example/
│       │   ├── config/            # Configuration and browser setup
│       │   │   ├── TestConfig.java
│       │   │   └── BrowserFactory.java
│       │   ├── core/              # Framework foundation
│       │   │   ├── DriverManager.java
│       │   │   └── BaseTest.java
│       │   ├── pages/             # Page Object Model
│       │   │   ├── BasePage.java
│       │   │   ├── theinternet/
│       │   │   │   ├── LoginPage.java
│       │   │   │   └── SecureAreaPage.java
│       │   │   └── saucedemo/
│       │   │       ├── LoginPage.java
│       │   │       └── InventoryPage.java
│       │   └── tests/             # E2E test classes
│       │       ├── TheInternetLoginTest.java
│       │       └── SauceDemoLoginTest.java
│       └── resources/
│           ├── config.properties  # Runtime configuration
│           ├── testng.xml         # TestNG suite (parallel config)
│           └── allure.properties  # Allure results directory
├── build.gradle.kts               # Build configuration and dependencies
├── settings.gradle.kts            # Project settings
├── gradlew / gradlew.bat          # Gradle wrapper scripts
└── .gitignore                     # Ignored files (build output, reports)""")

    # 4. Architecture
    add_heading(doc, "4. Architecture & Component Connections", 1)
    add_para(doc, "The framework is organized in layers. Data and control flow from configuration through driver management, page objects, and tests.", bold=False)

    add_heading(doc, "4.1 Execution Flow (Single Test)", 2)
    add_numbered(doc, "Gradle runs the test task, which invokes TestNG using testng.xml")
    add_numbered(doc, "TestNG discovers test classes in org.example.tests package")
    add_numbered(doc, "BaseTest.@BeforeMethod calls DriverManager.initDriver()")
    add_numbered(doc, "DriverManager calls BrowserFactory.createDriver()")
    add_numbered(doc, "BrowserFactory reads TestConfig (config.properties + env vars) to build WebDriver")
    add_numbered(doc, "Selenium Manager auto-downloads the correct browser driver")
    add_numbered(doc, "Test method creates Page Objects and performs actions")
    add_numbered(doc, "Page Objects use DriverManager.getDriver() and explicit waits from BasePage")
    add_numbered(doc, "Allure captures @Step annotations and test metadata")
    add_numbered(doc, "BaseTest.@AfterMethod quits driver; on failure, attaches screenshot + page source")
    add_numbered(doc, "Allure results written to build/allure-results/")

    add_heading(doc, "4.2 Component Dependency Map", 2)
    add_table(doc, ["Component", "Depends On", "Used By"], [
        ["TestConfig", "config.properties, env vars", "BrowserFactory, BasePage, Page Objects"],
        ["BrowserFactory", "TestConfig, Selenium", "DriverManager"],
        ["DriverManager", "BrowserFactory", "BaseTest, BasePage"],
        ["BaseTest", "DriverManager, TestConfig, Allure", "All test classes"],
        ["BasePage", "DriverManager, TestConfig", "All page objects"],
        ["Page Objects", "BasePage, TestConfig", "Test classes"],
        ["Test Classes", "BaseTest, Page Objects, AssertJ", "TestNG suite"],
        ["testng.xml", "Test classes", "Gradle test task"],
        ["e2e-tests.yml", "Gradle, testng.xml", "GitHub Actions"],
    ])

    add_heading(doc, "4.3 Layer Responsibilities", 2)
    add_table(doc, ["Layer", "Package", "Responsibility"], [
        ["Config", "org.example.config", "Load settings, create browser instances"],
        ["Core", "org.example.core", "Driver lifecycle, test setup/teardown, failure artifacts"],
        ["Pages", "org.example.pages", "Encapsulate page locators and user actions"],
        ["Tests", "org.example.tests", "Define test scenarios and assertions"],
        ["Resources", "src/test/resources", "External config files and TestNG suite"],
        ["CI/CD", ".github/workflows", "Automated test execution on GitHub"],
    ])

    # 5. Configuration
    add_heading(doc, "5. Configuration Reference", 1)

    add_heading(doc, "5.1 config.properties", 2)
    add_code_block(doc, """browser=chrome
headless=false
base.url.theinternet=https://the-internet.herokuapp.com
base.url.saucedemo=https://www.saucedemo.com
implicit.wait.seconds=0
explicit.wait.seconds=10""")
    add_table(doc, ["Property", "Default", "Description"], [
        ["browser", "chrome", "Browser to use: chrome, firefox, or edge"],
        ["headless", "false", "Run browser without visible UI (true/false)"],
        ["base.url.theinternet", "herokuapp URL", "Base URL for The Internet demo site"],
        ["base.url.saucedemo", "saucedemo.com", "Base URL for Sauce Demo site"],
        ["implicit.wait.seconds", "0", "Global implicit wait (0 = disabled, use explicit waits)"],
        ["explicit.wait.seconds", "10", "Default explicit wait timeout for page actions"],
    ])

    add_heading(doc, "5.2 Environment Variable Overrides", 2)
    add_para(doc, "Environment variables take priority over config.properties. This is how CI overrides settings:")
    add_table(doc, ["Env Variable", "CI Value", "Purpose"], [
        ["HEADLESS", "true", "Force headless mode in CI"],
        ["BROWSER", "chrome", "Force Chrome in CI"],
    ])

    add_heading(doc, "5.3 testng.xml", 2)
    add_code_block(doc, """<suite name="E2E Suite" parallel="methods" thread-count="2">
    <test name="Demo Tests">
        <packages>
            <package name="org.example.tests"/>
        </packages>
    </test>
</suite>""")
    add_bullet(doc, "parallel=\"methods\" — runs test methods concurrently")
    add_bullet(doc, "thread-count=\"2\" — uses 2 parallel threads")
    add_bullet(doc, "Requires ThreadLocal WebDriver in DriverManager for thread safety")

    add_heading(doc, "5.4 allure.properties", 2)
    add_code_block(doc, "allure.results.directory=build/allure-results")

    # 6. Running Tests
    add_heading(doc, "6. Running Tests (Headed & Headless)", 1)

    add_heading(doc, "6.1 Prerequisites", 2)
    add_bullet(doc, "Java 17 or higher installed (JAVA_HOME set)")
    add_bullet(doc, "Chrome, Firefox, or Edge browser installed locally")
    add_bullet(doc, "No manual chromedriver/geckodriver needed — Selenium Manager handles it")

    add_heading(doc, "6.2 Headed Mode (Visible Browser) — Default", 2)
    add_para(doc, "Windows (PowerShell):", bold=True)
    add_code_block(doc, """cd D:\\Code\\web-automation-selenium\\selenium
.\\gradlew.bat clean test""")
    add_para(doc, "Windows (CMD):", bold=True)
    add_code_block(doc, """cd D:\\Code\\web-automation-selenium\\selenium
gradlew.bat clean test""")
    add_para(doc, "Linux / macOS:", bold=True)
    add_code_block(doc, """cd /path/to/selenium
chmod +x gradlew
./gradlew clean test""")
    add_para(doc, "Uses headless=false from config.properties. Browser window will be visible.")

    add_heading(doc, "6.3 Headless Mode (No Visible Browser)", 2)
    add_para(doc, "Windows (PowerShell):", bold=True)
    add_code_block(doc, """$env:HEADLESS='true'
$env:BROWSER='chrome'
.\\gradlew.bat clean test""")
    add_para(doc, "Windows (CMD):", bold=True)
    add_code_block(doc, """set HEADLESS=true
set BROWSER=chrome
gradlew.bat clean test""")
    add_para(doc, "Linux / macOS:", bold=True)
    add_code_block(doc, """HEADLESS=true BROWSER=chrome ./gradlew clean test""")

    add_heading(doc, "6.4 Run a Single Test Class", 2)
    add_code_block(doc, """.\\gradlew.bat test --tests "org.example.tests.TheInternetLoginTest"
.\\gradlew.bat test --tests "org.example.tests.SauceDemoLoginTest\"""")

    add_heading(doc, "6.5 Run with Different Browsers", 2)
    add_code_block(doc, """# Firefox headed
$env:BROWSER='firefox'; .\\gradlew.bat test

# Edge headless
$env:BROWSER='edge'; $env:HEADLESS='true'; .\\gradlew.bat test""")

    add_heading(doc, "6.6 Run from IntelliJ IDEA", 2)
    add_numbered(doc, "Open the project in IntelliJ IDEA")
    add_numbered(doc, "Right-click testng.xml → Run, OR right-click a test class → Run")
    add_numbered(doc, "To run headless from IDE, add environment variables in Run Configuration: HEADLESS=true")

    # 7. Allure Reporting
    add_heading(doc, "7. Allure Reporting", 1)

    add_heading(doc, "7.1 Generate HTML Report", 2)
    add_code_block(doc, """.\\gradlew.bat clean test allureReport""")
    add_para(doc, "Report output: build/reports/allure-report/allureReport/index.html")

    add_heading(doc, "7.2 Open Report in Browser (Live Server)", 2)
    add_code_block(doc, """.\\gradlew.bat allureServe""")
    add_para(doc, "Starts a local server and opens the report automatically.")

    add_heading(doc, "7.3 What Allure Captures", 2)
    add_bullet(doc, "@Epic, @Feature, @Story, @Severity — test categorization")
    add_bullet(doc, "@Step — page object action steps (visible in report timeline)")
    add_bullet(doc, "Browser and headless parameters on each test")
    add_bullet(doc, "Screenshot + page source on test failure")
    add_bullet(doc, "Raw JSON results in build/allure-results/")

    add_heading(doc, "7.4 Report Output Locations", 2)
    add_table(doc, ["Output", "Path", "Description"], [
        ["Raw results", "build/allure-results/", "JSON files used to generate report"],
        ["HTML report", "build/reports/allure-report/allureReport/", "Interactive HTML dashboard"],
        ["Gradle test report", "build/reports/tests/test/", "Standard Gradle test HTML report"],
    ])

    # 8. CI/CD
    add_heading(doc, "8. CI/CD Pipeline (GitHub Actions)", 1)

    add_heading(doc, "8.1 Workflow File", 2)
    add_para(doc, "Location: .github/workflows/e2e-tests.yml")

    add_heading(doc, "8.2 Pipeline Steps", 2)
    add_table(doc, ["Step", "Action", "Details"], [
        ["1. Checkout", "actions/checkout@v4", "Clone repository code"],
        ["2. Setup Java", "actions/setup-java@v4", "Install Temurin JDK 17, cache Gradle deps"],
        ["3. chmod +x gradlew", "Shell command", "Fix executable bit (Windows commits lack it)"],
        ["4. Run E2E tests", "./gradlew clean test", "HEADLESS=true, BROWSER=chrome, Ubuntu runner"],
        ["5. Generate Allure", "./gradlew allureReport", "Runs even if tests fail (if: always())"],
        ["6. Upload artifacts", "actions/upload-artifact@v4", "Upload allure-results and allure-report"],
    ])

    add_heading(doc, "8.3 CI Environment", 2)
    add_table(doc, ["Setting", "Value"], [
        ["Runner OS", "ubuntu-latest"],
        ["Java", "17 (Temurin)"],
        ["Browser", "Chrome (pre-installed on Ubuntu runner)"],
        ["Mode", "Headless"],
        ["Driver management", "Selenium Manager (automatic)"],
    ])

    add_heading(doc, "8.4 Pipeline Flow Diagram", 2)
    add_code_block(doc, """Push/PR to main or master
        │
        ▼
  GitHub Actions triggered
        │
        ▼
  Checkout code → Setup Java 17 → chmod gradlew
        │
        ▼
  ./gradlew clean test (headless Chrome)
        │
        ├── Tests PASS → continue
        └── Tests FAIL → still continue (if: always)
        │
        ▼
  ./gradlew allureReport
        │
        ▼
  Upload artifacts (allure-results, allure-report)
        │
        ▼
  Workflow completes (green or red based on test result)""")

    # 9. Trigger CI/CD
    add_heading(doc, "9. How to Trigger CI/CD", 1)

    add_heading(doc, "9.1 Automatic Triggers", 2)
    add_bullet(doc, "Push to master or main branch → workflow runs automatically")
    add_bullet(doc, "Open or update a Pull Request → workflow runs automatically")

    add_heading(doc, "9.2 Manual Trigger (Push Any Change)", 2)
    add_code_block(doc, """git add .
git commit -m "Your change description"
git push origin master""")

    add_heading(doc, "9.3 Re-run Failed Workflow", 2)
    add_numbered(doc, "Go to https://github.com/Gourab-Pal/selenium/actions")
    add_numbered(doc, "Click the failed workflow run")
    add_numbered(doc, "Click 'Re-run all jobs' button")

    add_heading(doc, "9.4 View CI Results", 2)
    add_numbered(doc, "Actions tab: https://github.com/Gourab-Pal/selenium/actions")
    add_numbered(doc, "Click a workflow run to see step-by-step logs")
    add_numbered(doc, "Scroll to 'Artifacts' section at the bottom to download Allure reports")

    add_heading(doc, "9.5 Download Allure Report from CI", 2)
    add_numbered(doc, "Open a completed workflow run")
    add_numbered(doc, "Download the 'allure-report' artifact (ZIP)")
    add_numbered(doc, "Extract and open allureReport/index.html in a browser")

    # 10. Test Suite Details
    add_heading(doc, "10. Test Suite Details", 1)

    add_heading(doc, "10.1 TheInternetLoginTest", 2)
    add_table(doc, ["Field", "Value"], [
        ["Class", "org.example.tests.TheInternetLoginTest"],
        ["Site", "https://the-internet.herokuapp.com"],
        ["Flow", "Login → Verify secure area"],
        ["Credentials", "tomsmith / SuperSecretPassword!"],
        ["Assertions", "Heading = 'Secure Area', success banner contains login message"],
        ["Allure", "Epic: Demo Sites, Feature: The Internet Login, Severity: CRITICAL"],
    ])

    add_heading(doc, "10.2 SauceDemoLoginTest", 2)
    add_table(doc, ["Field", "Value"], [
        ["Class", "org.example.tests.SauceDemoLoginTest"],
        ["Site", "https://www.saucedemo.com"],
        ["Flow", "Login → Verify inventory page"],
        ["Credentials", "standard_user / secret_sauce"],
        ["Assertions", "Inventory page displayed, title = 'Products'"],
        ["Allure", "Epic: Demo Sites, Feature: Sauce Demo Login, Severity: CRITICAL"],
    ])

    # 11. Page Object Model
    add_heading(doc, "11. Page Object Model (POM)", 1)
    add_para(doc, "Each web page is represented by a Java class that encapsulates:")
    add_bullet(doc, "Locators (By.id, By.cssSelector) — private static final fields")
    add_bullet(doc, "Actions (click, type, navigate) — public methods with @Step for Allure")
    add_bullet(doc, "Fluent API — methods return this or the next page object for chaining")

    add_heading(doc, "11.1 Example: Login Flow", 2)
    add_code_block(doc, """// In test class:
SecureAreaPage page = new LoginPage().loginAs("tomsmith", "SuperSecretPassword!");

// loginAs() chains:
// open() → enterUsername() → enterPassword() → submitLogin() → SecureAreaPage""")

    add_heading(doc, "11.2 BasePage Utilities", 2)
    add_table(doc, ["Method", "Purpose"], [
        ["open(url)", "Navigate to URL"],
        ["waitForVisible(locator)", "Explicit wait until element is visible"],
        ["waitForClickable(locator)", "Explicit wait until element is clickable"],
        ["click(locator)", "Wait + click"],
        ["type(locator, text)", "Wait + clear + sendKeys"],
        ["getText(locator)", "Wait + get text"],
    ])

    # 12. Adding New Tests
    add_heading(doc, "12. Adding New Tests", 1)
    add_numbered(doc, "Create page object(s) in src/test/java/org/example/pages/<sitename>/")
    add_numbered(doc, "Extend BasePage, define locators and @Step methods")
    add_numbered(doc, "Add base URL to config.properties if needed")
    add_numbered(doc, "Create test class in src/test/java/org/example/tests/ extending BaseTest")
    add_numbered(doc, "Add @Test, Allure annotations, and AssertJ assertions")
    add_numbered(doc, "TestNG auto-discovers via testng.xml package scan (org.example.tests)")
    add_numbered(doc, "Run locally: .\\gradlew.bat clean test")
    add_numbered(doc, "Push to GitHub to trigger CI")

    add_heading(doc, "12.1 Template: New Test Class", 2)
    add_code_block(doc, """package org.example.tests;

import org.example.core.BaseTest;
import org.testng.annotations.Test;
import static org.assertj.core.api.Assertions.assertThat;

public class MyNewTest extends BaseTest {

    @Test(description = "Description of what this test verifies")
    public void myTestMethod() {
        // Use page objects here
        // assertThat(...).isEqualTo(...);
    }
}""")

    # 13. Troubleshooting
    add_heading(doc, "13. Troubleshooting", 1)
    add_table(doc, ["Problem", "Cause", "Solution"], [
        ["CI fails instantly on gradlew", "gradlew not executable on Linux", "chmod +x gradlew (already in workflow)"],
        ["WebDriverException: driver not found", "Browser not installed", "Install Chrome/Firefox/Edge locally"],
        ["Test timeout", "Element not found in time", "Increase explicit.wait.seconds in config.properties"],
        ["Allure report empty", "Tests not run before report", "Run: gradlew clean test allureReport"],
        ["Parallel test failures", "Shared state between tests", "Ensure DriverManager uses ThreadLocal (already done)"],
        ["HEADLESS not working", "Env var not set correctly", "Use $env:HEADLESS='true' in PowerShell"],
        ["Permission denied on Linux", "gradlew mode 100644", "Run chmod +x gradlew or use workflow fix"],
    ])

    # 14. Quick Command Reference
    add_heading(doc, "14. Quick Command Reference", 1)
    add_table(doc, ["Task", "Command (Windows PowerShell)"], [
        ["Run all tests (headed)", ".\\gradlew.bat clean test"],
        ["Run all tests (headless)", "$env:HEADLESS='true'; .\\gradlew.bat clean test"],
        ["Generate Allure report", ".\\gradlew.bat allureReport"],
        ["Test + report", ".\\gradlew.bat clean test allureReport"],
        ["Open report in browser", ".\\gradlew.bat allureServe"],
        ["Run single test", ".\\gradlew.bat test --tests \"org.example.tests.TheInternetLoginTest\""],
        ["Run with Firefox", "$env:BROWSER='firefox'; .\\gradlew.bat test"],
        ["Push and trigger CI", "git push origin master"],
        ["View CI runs", "https://github.com/Gourab-Pal/selenium/actions"],
    ])

    # Footer
    doc.add_page_break()
    add_heading(doc, "Appendix: Git Repository Info", 1)
    add_table(doc, ["Item", "Value"], [
        ["Remote URL", "https://github.com/Gourab-Pal/selenium.git"],
        ["Default Branch", "master"],
        ["CI Workflow", ".github/workflows/e2e-tests.yml"],
        ["CI Status URL", "https://github.com/Gourab-Pal/selenium/actions"],
        ["Latest Successful Run", "https://github.com/Gourab-Pal/selenium/actions/runs/27117702335"],
    ])

    return doc


if __name__ == "__main__":
    output_path = r"D:\Code\web-automation-selenium\selenium\docs\Selenium-E2E-Framework-Documentation.docx"
    import os
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    document = build_document()
    document.save(output_path)
    print(f"Documentation saved to: {output_path}")
